#!/usr/bin/env python3
"""
Script to convert Quiz Master Project Report from Markdown to DOCX format
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re

def create_docx_from_markdown(markdown_file, output_file):
    """Convert markdown file to DOCX with proper formatting"""
    
    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set up document styles
    setup_document_styles(doc)
    
    # Split content into lines
    lines = md_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            add_heading(doc, text, level)
        
        # Handle code blocks
        elif line.startswith('```'):
            # Find the end of code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, '\n'.join(code_lines))
        
        # Handle regular paragraphs
        elif line and not line.startswith('---'):
            # Check if it's a list item
            if line.startswith('- ') or line.startswith('* '):
                add_list_item(doc, line[2:])
            elif line.startswith('1. '):
                add_numbered_item(doc, line[3:])
            else:
                add_paragraph(doc, line)
        
        i += 1
    
    # Save the document
    doc.save(output_file)
    print(f"✅ Successfully converted {markdown_file} to {output_file}")

def setup_document_styles(doc):
    """Set up document styles for consistent formatting"""
    
    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    h1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    h2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(6)
    
    # Heading 3 style
    h3_style = doc.styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'Arial'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(4)
    
    # Code style
    code_style = doc.styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

def add_heading(doc, text, level):
    """Add heading with appropriate level"""
    if level == 1:
        p = doc.add_paragraph(text, style='CustomTitle')
    elif level == 2:
        p = doc.add_paragraph(text, style='CustomHeading1')
    elif level == 3:
        p = doc.add_paragraph(text, style='CustomHeading2')
    else:
        p = doc.add_paragraph(text, style='CustomHeading3')

def add_paragraph(doc, text):
    """Add regular paragraph"""
    if text:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

def add_list_item(doc, text):
    """Add bullet list item"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

def add_numbered_item(doc, text):
    """Add numbered list item"""
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(3)

def add_code_block(doc, code_text):
    """Add code block with monospace font"""
    # Split code into lines and add each line
    lines = code_text.split('\n')
    for line in lines:
        if line.strip():
            p = doc.add_paragraph(line, style='CustomCode')
        else:
            p = doc.add_paragraph('', style='CustomCode')

def add_mermaid_diagram_placeholder(doc):
    """Add placeholder for Mermaid diagram"""
    p = doc.add_paragraph("ER Diagram - Database Schema", style='CustomHeading2')
    p = doc.add_paragraph("Note: The ER diagram is included as a Mermaid diagram in the original markdown. For the DOCX version, please refer to the markdown file or generate the diagram separately using a Mermaid renderer.")
    p.paragraph_format.italic = True

if __name__ == "__main__":
    input_file = "Quiz_Master_Project_Report.md"
    output_file = "Quiz_Master_Project_Report.docx"
    
    try:
        create_docx_from_markdown(input_file, output_file)
        print(f"📄 Document saved as: {output_file}")
        print("📝 Note: The ER diagram (Mermaid format) will need to be converted separately to an image format for the DOCX.")
    except Exception as e:
        print(f"❌ Error converting file: {e}")
        print("💡 Make sure you have the 'python-docx' package installed:")
        print("   pip install python-docx")