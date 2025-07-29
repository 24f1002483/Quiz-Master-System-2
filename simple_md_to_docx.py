#!/usr/bin/env python3
"""
Simple script to convert Quiz Master Project Report from Markdown to DOCX format
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def convert_markdown_to_docx(markdown_file, output_file):
    """Convert markdown file to DOCX with proper formatting"""
    
    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set up document styles
    setup_styles(doc)
    
    # Process content
    process_content(doc, content)
    
    # Save the document
    doc.save(output_file)
    print(f"✅ Successfully converted {markdown_file} to {output_file}")

def setup_styles(doc):
    """Set up document styles"""
    
    # Title style
    title_style = doc.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    h1_style = doc.styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    h2_style = doc.styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(6)
    
    # Heading 3 style
    h3_style = doc.styles.add_style('Heading3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'Arial'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(4)
    
    # Code style
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.3)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

def process_content(doc, content):
    """Process the markdown content and add to document"""
    
    # Split into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines and horizontal rules
        if not line or line == '---':
            i += 1
            continue
        
        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            add_heading(doc, text, level)
        
        # Handle code blocks
        elif line.startswith('```'):
            # Find the end of code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                add_code_block(doc, '\n'.join(code_lines))
        
        # Handle list items
        elif line.startswith('- ') or line.startswith('* '):
            add_list_item(doc, line[2:])
        
        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            add_numbered_item(doc, text)
        
        # Handle regular paragraphs
        else:
            add_paragraph(doc, line)
        
        i += 1

def add_heading(doc, text, level):
    """Add heading with appropriate level"""
    if level == 1:
        doc.add_paragraph(text, style='Title')
    elif level == 2:
        doc.add_paragraph(text, style='Heading1')
    elif level == 3:
        doc.add_paragraph(text, style='Heading2')
    else:
        doc.add_paragraph(text, style='Heading3')

def add_paragraph(doc, text):
    """Add regular paragraph"""
    if text:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

def add_list_item(doc, text):
    """Add bullet list item"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

def add_numbered_item(doc, text):
    """Add numbered list item"""
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(3)

def add_code_block(doc, code_text):
    """Add code block with monospace font"""
    # Check if it's a Mermaid diagram
    if 'erDiagram' in code_text:
        # Add ER Diagram section
        doc.add_paragraph("ER Diagram - Database Schema", style='Heading2')
        doc.add_paragraph("The following ER diagram shows the database schema with all tables and their relationships:")
        
        # Add the Mermaid code as a code block
        p = doc.add_paragraph("Mermaid ER Diagram Code:", style='Heading3')
        p = doc.add_paragraph(code_text, style='Code')
        
        # Add note about diagram
        note = doc.add_paragraph("Note: This is a Mermaid diagram code. To view the actual diagram, you can:")
        note.paragraph_format.italic = True
        doc.add_paragraph("1. Use an online Mermaid renderer (e.g., mermaid.live)", style='List Number')
        doc.add_paragraph("2. Copy the code above and paste it into the renderer", style='List Number')
        doc.add_paragraph("3. Or refer to the original markdown file for the rendered diagram", style='List Number')
    else:
        # Regular code block
        p = doc.add_paragraph(code_text, style='Code')

if __name__ == "__main__":
    input_file = "Quiz_Master_Project_Report.md"
    output_file = "Quiz_Master_Project_Report.docx"
    
    try:
        convert_markdown_to_docx(input_file, output_file)
        print(f"📄 Document saved as: {output_file}")
        print("📝 Note: The ER diagram is included as Mermaid code. You can render it using mermaid.live")
    except ImportError:
        print("❌ Error: python-docx package not found")
        print("💡 Install it using: pip install python-docx")
    except Exception as e:
        print(f"❌ Error converting file: {e}")